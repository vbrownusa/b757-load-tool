import streamlit as st
import pandas as pd
from docx import Document

# -------------------------
# BAG AWU LOOKUP (RANGE BASED)
# -------------------------

BAG_TABLE = [
    (0, 10, 0.0),
    (11, 20, 640.0),
    (21, 30, 960.0),
    (31, 40, 1280.0),
    (41, 50, 1600.0),
    (51, 60, 1920.0),
    (61, 70, 2240.0),
    (71, 80, 2560.0),
]

# -------------------------
# CARGO AWU (YOUR RULE)
# -------------------------

def cargo_awu_by_rule(table, weight):
    if weight <= 0:
        return 0.0

    weight = float(weight)

    awu_values = sorted(table.values())

    for awu in awu_values:
        # ignore last digits → convert to hundreds
        equiv_lbs = int(awu) - (int(awu) % 100)

        if equiv_lbs >= weight:
            return awu

    return awu_values[-1]


def bag_awu(count):
    count = int(count)

    for low, high, awu in BAG_TABLE:
        if low <= count <= high:
            return awu

    st.error(f"Bag count out of range: {count}")
    return 0.0



# -------------------------
# LOAD PASSENGER DATA
# -------------------------


import streamlit as st
import pandas as pd
from docx import Document

@st.cache_data
def load_pax_data():
    df = pd.read_csv("pax_data.csv")

    df.columns = df.columns.str.strip().str.lower()
    df = df.dropna()

    df["zone"] = df["zone"].astype(str).str.strip().str.upper()
    df["pax"] = df["pax"].astype(int)

    return df

# 👉 THIS LINE IS CRITICAL
pax_df = load_pax_data()


# -------------------------
# ✅ ADD TABLES HERE (TOP LEVEL)
# -------------------------

ZONE_A = { ... }
ZONE_B = { ... }
ZONE_C = { ... }

# -------------------------
# FUNCTIONS
# -------------------------

def pax_awu(df, zone, pax, season):
    if pax == 0:
        return 0.0

    zone = str(zone).strip().upper()
    pax = int(pax)

    row = df[(df["zone"] == zone) & (df["pax"] == pax)]

    if row.empty:
        st.error(f"Missing AWU data for Zone {zone}, Pax {pax}")
        return 0.0

    return float(row.iloc[0][season])
    
def generate_release(data):
    ...
    
# -------------------------
# UI (BOTTOM)
# -------------------------

st.title("B757 Weight and Balance Tool")

# -------------------------
# TEST FUNCTION (CLEAN)
# -------------------------

def generate_release(data):
    doc = Document()

    def add(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold

    add("B757 Weight and Balance Key", True)

    add("PASSENGERS", True)
    add(f"Zone A: {data['zoneA']}")
    add(f"Zone B: {data['zoneB']}")
    add(f"Zone C: {data['zoneC']}")

    add("BAGGAGE (DOMESTIC)", True)
    add(f"Bin 1: {data['bag1']}")
    add(f"Bin 2: {data['bag2']}")
    add(f"Bin 3: {data['bag3']}")
    add(f"Bin 4: {data['bag4']}")

    add("REVENUE CARGO", True)
    add(f"Bin 1: {data['cargo_vals'][0]}")
    add(f"Bin 2: {data['cargo_vals'][1]}")
    add(f"Bin 3: {data['cargo_vals'][2]}")
    add(f"Bin 4: {data['cargo_vals'][3]}")

    file = "release.docx"
    doc.save(file)

    return file


# -------------------------
# SIMPLE TEST UI
# -------------------------

season = st.radio("Season", ["summer", "winter"], horizontal=True)

col1, col2, col3 = st.columns(3)


with col1:
    a = int(st.number_input("Zone A", min_value=0, max_value=54, step=1))
    b = int(st.number_input("Zone B", min_value=0, max_value=80, step=1))
    c = int(st.number_input("Zone C", min_value=0, max_value=84, step=1))

with col2:
    b1 = st.number_input("Bin 1 Bags", 0)
    b2 = st.number_input("Bin 2 Bags", 0)
    b3 = st.number_input("Bin 3 Bags", 0)
    b4 = st.number_input("Bin 4 Bags", 0)

with col3:
    c1 = st.number_input("Cargo Bin 1", min_value=0, value=0)
    c2 = st.number_input("Cargo Bin 2", min_value=0, value=0)
    c3 = st.number_input("Cargo Bin 3", min_value=0, value=0)
    c4 = st.number_input("Cargo Bin 4", min_value=0, value=0)

rf = st.number_input("Ramp Fuel", 0)
tf = st.number_input("Taxi Fuel", 0)


# -------------------------
# CALCULATIONS (PUT IT HERE)
# -------------------------

za = pax_awu(pax_df, "A", a, season)
zb = pax_awu(pax_df, "B", b, season)
zc = pax_awu(pax_df, "C", c, season)

bag_awus = [
    bag_awu(b1),
    bag_awu(b2),
    bag_awu(b3),
    bag_awu(b4),
]

cargo_awus = [
    cargo_awu_by_rule(CARGO_BIN1, c1),
    cargo_awu_by_rule(CARGO_BIN2, c2),
    cargo_awu_by_rule(CARGO_BIN3, c3),
    cargo_awu_by_rule(CARGO_BIN4, c4),
]

tof = rf - tf
fuel_awu_val = fuel_awu(tof)

zfw = BOW + za + zb + zc + sum(bag_awus) + sum(cargo_awus)

BOW = 129621.4

za = pax_awu(pax_df, "A", a, season)
zb = pax_awu(pax_df, "B", b, season)
zc = pax_awu(pax_df, "C", c, season)

bags = [b1 * 32, b2 * 32, b3 * 32, b4 * 32]

cargo = [c1, c2, c3, c4]

zfw = BOW + za + zb + zc + sum(bags) + sum(cargo)

st.write("ZFW:", round(zfw,1))
